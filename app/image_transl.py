import json
import os
from concurrent.futures import ThreadPoolExecutor
import cv2
import numpy as np
import pytesseract
from PIL import Image, ImageOps
from docx import Document
import concurrent.futures
from app.translation import generate
from functools import wraps
import threading
from win32com import client as wc


def preprocess_image(image_path, cutoff=2, block_size=11, C=2, h=3, templateWindowSize=7, searchWindowSize=21):
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"The image {image_path} does not exist.")

    image = Image.open(image_path).convert('L')  # 灰度化
    image = ImageOps.autocontrast(image, cutoff=cutoff)  # 降低对比度调整强度
    image = np.array(image)
    # 自适应阈值
    thresh = cv2.adaptiveThreshold(image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block_size, C)
    # 去噪
    denoised = cv2.fastNlMeansDenoising(thresh, h=h, templateWindowSize=templateWindowSize,
                                        searchWindowSize=searchWindowSize)
    # 边缘增强
    kernel = np.ones((1, 1), np.uint8)
    enhanced = cv2.morphologyEx(denoised, cv2.MORPH_CLOSE, kernel)
    # 转换回RGB以便保存
    enhanced_rgb = cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR)
    cv2.imwrite(r'C:\Users\yxy\PycharmProjects\pythonProject\app\resources\preprocess.png', enhanced_rgb)
    return cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR)


# oem：OCR 引擎模式 psm页面分割模式
def image_text(image_path, lang='chi_sim+eng', oem=3, psm=6):
    try:
        processed_img = preprocess_image(image_path)
        pil_img = Image.fromarray(processed_img)
        custom_config = f'--oem {oem} --psm {psm}'
        text = pytesseract.image_to_string(pil_img, lang=lang, config=custom_config)
        return text
    except Exception as e:
        print(f"An error: {e}")
        return None


def overlay_text_on_image(image_path, translated_text, output_path=None):
    try:
        # 读取原始图片
        img = cv2.imread(image_path)
        if img is None:
            raise ValueError("无法读取图片文件")

        # 设置文本参数
        font = cv2.FONT_HERSHEY_SIMPLEX
        font_scale = 1.2
        thickness = 2
        color = (0, 0, 255)  # 红色文本
        line_type = cv2.LINE_AA

        # 计算文本位置(居中)
        (text_width, text_height), _ = cv2.getTextSize(translated_text, font, font_scale, thickness)
        img_height, img_width = img.shape[:2]
        x = int((img_width - text_width) / 2)
        y = int((img_height + text_height) / 2)

        # 添加文本背景(提高可读性)
        cv2.rectangle(img,
                      (x - 5, y - text_height - 5),
                      (x + text_width + 5, y + 5),
                      (255, 255, 255), -1)

        # 添加文本
        cv2.putText(img, translated_text, (x, y), font,
                    font_scale, color, thickness, line_type)

        # 保存结果
        if output_path is None:
            output_path = os.path.splitext(image_path)[0] + '_translated.png'
        cv2.imwrite(output_path, img)
        return output_path
    except Exception as e:
        print(f"图片文本覆盖失败: {e}")
        return None


def ocr_transl(sourceLang, targetLang, imgText, image_path=None):
    generate(sourceLang, targetLang, imgText)
    result = {
        'source': imgText,
        'first': '',
        'reflect': '',
        'improve': ''
    }
    for chunk in generate(sourceLang, targetLang, imgText):
        try:
            data = json.loads(chunk.strip().replace('data: ', ''))
            if data.get('stage') == 'first':
                result['first'] += data.get('chunk', '')
            elif data.get('stage') == 'reflect':
                result['reflect'] += data.get('chunk', '')
            elif data.get('stage') == 'improve':
                result['improve'] += data.get('chunk', '')
        except Exception as e:
            print(f"Error processing chunk: {e}")
    if image_path and os.path.exists(image_path):
        output_path = overlay_text_on_image(image_path, result['improve'])
        if output_path:
            print(f"翻译结果已覆盖到图片: {output_path}")
    return result['improve']


def convert_doc_to_docx(path):
    word = wc.Dispatch("Word.Application")
    word.Visible = 0
    word.DisplayAlerts = 0
    doc = word.Documents.Open(path)
    docx_path = os.path.splitext(path)[0] + '.docx'
    doc.SaveAs(docx_path, 12, False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()
    os.remove(path)
    return docx_path


def thread_tracker(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        thread_id = threading.get_ident()
        print(f"线程 {thread_id} 开始处理任务")
        try:
            result = func(*args, **kwargs)
            print(f"线程 {thread_id} 完成任务")
            return result
        except Exception as e:
            print(f"线程 {thread_id} 处理失败: {str(e)}")
            raise

    return wrapper


@thread_tracker
def translate_run(args):
    run, sourceLang, targetLang = args
    if not run.text.strip():
        return run, run.text  # 返回原始run和空文本
    try:
        translated = ocr_transl(sourceLang, targetLang, run.text)
        return run, translated
    except Exception as e:
        print(f"翻译出错: {e}")
        return run, run.text


@thread_tracker
def translate_cell(args):
    """包装cell翻译任务便于线程池调用"""
    cell, sourceLang, targetLang = args
    if not cell.text.strip():
        return cell, cell.text

    try:
        translated = ocr_transl(sourceLang, targetLang, cell.text)
        return cell, translated if translated.strip() else cell.text
    except Exception as e:
        print(f"表格翻译出错: {e}")
        return cell, cell.text


@thread_tracker
def translate_merged_run(args):
    run, sourceLang, targetLang, text = args
    try:
        if all(c in '，。、；：‘’“”【】（）！？…—' for c in text):
            return run, text
        return run, ocr_transl(sourceLang, targetLang, text)
    except Exception as e:
        return run, text


def merge_runs(paragraph):
    merged_runs = []
    buffer = []
    for run in paragraph.runs:
        text = run.text
        # 如果是纯标点且缓冲区有内容
        if text.strip() and all(c in '，。、；：‘’“”【】（）！？…—' for c in text):
            if buffer:
                buffer.append((run, text))
            else:
                merged_runs.append((run, text))
        else:
            if buffer:
                # 合并缓冲区内容
                merged_text = ''.join(t for _, t in buffer)
                merged_runs.append((buffer[0][0], merged_text))
                buffer = []
            merged_runs.append((run, text))
    # 处理剩余的缓冲区内容
    if buffer:
        merged_text = ''.join(t for _, t in buffer)
        merged_runs.append((buffer[0][0], merged_text))
    return merged_runs


def translate_word(sourceLang, targetLang, input_path):
    if input_path.lower().endswith('.doc'):
        try:
            input_path = convert_doc_to_docx(input_path)
        except Exception as e:
            print(f"转换失败: {e}")
            return
    doc = Document(input_path)
    print(f"找到段落: {len(doc.paragraphs)}")
    new_doc = Document()

    # 收集所有翻译任务
    all_tasks = []
    para_map = {}
    table_map = {}

    # 段落任务收集
    for para_idx, para in enumerate(doc.paragraphs):
        para_map[para_idx] = {
            'original': para,
            'new_para': new_doc.add_paragraph(),
            'runs': []
        }
        merged_runs = merge_runs(para)
        for run, text in merged_runs:
            if text.strip():
                all_tasks.append(('run', para_idx, run, sourceLang, targetLang, text))

    # 表格任务收集
    for table_idx, table in enumerate(doc.tables):
        table_map[table_idx] = {
            'original': table,
            'new_table': new_doc.add_table(rows=len(table.rows), cols=len(table.columns)),
            'cells': []
        }
        for row_idx, row in enumerate(table.rows):
            table_map[table_idx]['cells'].append([])
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    all_tasks.append(('cell', table_idx, row_idx, cell_idx, cell, sourceLang, targetLang))

    # 并行处理翻译任务
    with ThreadPoolExecutor(max_workers=8) as executor:
        future_to_task = {}

        # 提交所有任务
        for task in all_tasks:
            if task[0] == 'run':
                future = executor.submit(translate_merged_run, task[2:])
            else:
                future = executor.submit(translate_cell, (task[4], task[5], task[6]))
            future_to_task[future] = task

        # 处理结果
        for future in concurrent.futures.as_completed(future_to_task):
            task = future_to_task[future]
            if task[0] == 'run':
                _, para_idx, run, _, _, _ = task
                _, translated = future.result()
                para_map[para_idx]['runs'].append((run, translated))
            else:
                _, table_idx, row_idx, cell_idx, cell, _, _ = task
                _, translated = future.result()
                # 直接填充表格内容
                new_cell = table_map[table_idx]['new_table'].cell(row_idx, cell_idx)
                new_cell.text = translated if translated.strip() else cell.text
                # 复制基础格式
                new_cell.vertical_alignment = cell.vertical_alignment

    # 重建段落格式
    for para_info in para_map.values():
        for original_run, translated in para_info['runs']:
            new_run = para_info['new_para'].add_run(translated)
            # 复制格式
            new_run.bold = original_run.bold
            new_run.italic = original_run.italic
            new_run.underline = original_run.underline
            new_run.font.name = original_run.font.name
    output_dir = r"C:\Users\yxy\PycharmProjects\pythonProject\app\img"
    original_filename = os.path.splitext(os.path.basename(input_path))[0]
    output_name = f"{original_filename}_output.docx"
    output_path = os.path.join(output_dir, output_name)
    new_doc.save(output_path)
    return output_name
